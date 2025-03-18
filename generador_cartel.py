import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import os

def obtener_dia_semana(fecha, idiomas):
    dias = {
        "Español": ["Lunes", "Martes", "Miércoles", "Jueves", "Viernes", "Sábado", "Domingo"],
        "Portugués": ["Segunda-Feira", "Terça-Feira", "Quarta-Feira", "Quinta-Feira", "Sexta-Feira", "Sábado", "Domingo"],
        "Inglés": ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]
    }
    try:
        fecha_dt = datetime.strptime(fecha, "%d/%m/%Y")
        dias_traducidos = [dias.get(idioma, dias["Español"])[fecha_dt.weekday()] for idioma in idiomas]
        return f"{' / '.join(dias_traducidos)} - {fecha}"
    except ValueError:
        return "Día inválido"

def generar_cartel(ciudad, fecha, actividad, hora_encuentro, punto_encuentro, desayuno, nombre_guia, op1, precio_op1, op2, precio_op2, idiomas):
    doc_path = "EJEMPLO CARTEL EMV.docx"
    if not os.path.exists(doc_path):
        return "Error: No se encuentra el archivo base. Asegúrate de que 'EJEMPLO CARTEL EMV.docx' está en el directorio."
    
    doc = Document(doc_path)
    
    fecha_formateada = obtener_dia_semana(fecha, idiomas)
    
    traducciones = {
        "Español": {"Bienvenidos": "¡Bienvenidos!", "Guía": "GUÍA", "Opcional": "Paseo opcional", "NoOpcionales": "No hay Excursiones Opcionales para el Día de Hoy", "Actividad": "Actividad", "Desayuno": "Desayuno", "Salida": "Salida", "PuntodeEncuentro": "Punto de Encuentro", "HoradeEncuentro": "Hora de Encuentro"},
        "Portugués": {"Bienvenidos": "Bem-Vindos!", "Guía": "GUIA", "Opcional": "Passeio opcional", "NoOpcionales": "Não há passeios opcionais para hoje", "Actividad": "Atividade", "Desayuno": "Café da Manhã", "Salida": "Saída", "PuntodeEncuentro": "Ponto de Encontro", "HoradeEncuentro": "Hora de Encontro"},
        "Inglés": {"Bienvenidos": "Welcome!", "Guía": "GUIDE", "Opcional": "Optional excursion", "NoOpcionales": "There are no optional excursions for today", "Actividad": "Activity", "Desayuno": "Breakfast", "Salida": "Departure", "PuntodeEncuentro": "Meeting Point", "HoradeEncuentro": "Meeting Hour"}
    }
    
    textos_traducidos = [traducciones.get(idioma, traducciones["Español"]) for idioma in idiomas]
    
    bienvenida = " / ".join([texto['Bienvenidos'] for texto in textos_traducidos])
    guia_traducido = " / ".join([texto['Guía'] for texto in textos_traducidos])
    actividad_traducida = " / ".join([texto['Actividad'] for texto in textos_traducidos]) + f" - {actividad}"
    desayuno_traducido = " / ".join([texto['Desayuno'] for texto in textos_traducidos]) + f": {desayuno}"
    no_opcionales_texto = " / ".join([texto['NoOpcionales'] for texto in textos_traducidos])
    punto_de_encuentro = " / ".join([texto['PuntodeEncuentro'] for texto in textos_traducidos])
    hora_de_encuentro = " / ".join([texto['HoradeEncuentro'] for texto in textos_traducidos])
    
    reemplazos = {
        "(BIENVENIDA)": bienvenida,
        "(CIUDAD)": f"{ciudad}",
        "📅": f"📅 {fecha_formateada}",
        "🥐": f"🥐 {desayuno_traducido}",
        "🚌": f"🚌 {actividad_traducida}",
        "⏰": f"⏰ {hora_de_encuentro}:\n{hora_encuentro}",
        "📍": f"📍 {punto_de_encuentro}:\n{punto_encuentro}",
        "🧑‍💼": f"🧑‍💼 {guia_traducido}: {nombre_guia}"
    }
    
    for p in doc.paragraphs:
        for key, value in reemplazos.items():
            if key in p.text:
                p.text = p.text.replace(key, value)
                for run in p.runs:
                    if key in ["(BIENVENIDA)", "(CIUDAD)"]:
                        run.font.name = "Neulis Sans Black"
                        run.font.size = Pt(18)
                        run.font.color.rgb = RGBColor(44, 66, 148)
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    elif key == "📅":
                        run.font.name = "Neulis Sans Black"
                        run.font.size = Pt(14)
                        run.font.color.rgb = RGBColor(44, 66, 148)
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    elif key == "🥐":
                        run.font.name = "Neulis Sans"
                        run.font.size = Pt(14)
                        run.font.color.rgb = RGBColor(44, 66, 148)
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    elif key == "🚌":
                        run.font.name = "Neulis Sans Black"
                        run.font.size = Pt(14)
                        run.font.color.rgb = RGBColor(44, 66, 148)
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    elif "⏰" in p.text:
                        run.font.name = "Neulis Sans Black"
                        run.font.size = Pt(16)
                        run.font.color.rgb = RGBColor(44, 66, 148)
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    else:
                        run.font.name = "Neulis Sans"
                        run.font.size = Pt(14)
                        run.font.color.rgb = RGBColor(44, 66, 148)
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        if "✨ Paseo opcional / Passeio opcional / Optional excursion" in p.text:
            if not op1 and not op2:
                opcional_run = p.add_run(f"\n{no_opcionales_texto}")
            else:
                if op1:
                    opcional_run = p.add_run(f"\n{op1} - 💰 {precio_op1}")
                if op2:
                    opcional_run = p.add_run(f"\n{op2} - 💰 {precio_op2}")
                opcional_run.font.name = "Neulis Sans"
                opcional_run.font.size = Pt(14)
                opcional_run.font.color.rgb = RGBColor(44, 66, 148)
    
    output_path = os.path.join(os.getcwd(), f"Cartel_{ciudad}_{'_'.join(idiomas)}.docx")
    doc.save(output_path)
    return output_path
st.title("Generador de Carteles para Pasajeros")

idiomas_disponibles = ["Español", "Portugués", "Inglés"]
idiomas_seleccionados = st.multiselect("Seleccione los idiomas:", idiomas_disponibles, default=["Español"])

if len(idiomas_seleccionados) == 0:
    st.warning("Debe seleccionar al menos un idioma para generar el cartel.")
else:
    ciudad = st.text_input("Ingrese la ciudad:")
    fecha = st.text_input("Ingrese la fecha (dd/mm/aaaa):")
    actividad = st.text_input("Ingrese el nombre de la actividad principal:")
    hora_encuentro = st.text_input("Ingrese la hora de encuentro:")
    punto_encuentro = st.text_input("Ingrese el punto de encuentro:")
    desayuno = st.text_input("Ingrese la hora del desayuno:")
    nombre_guia = st.text_input("Ingrese el nombre del guía:")
    op1 = st.text_input("Ingrese la Excursión Opcional 1 (Opcional):")
    precio_op1 = st.text_input("Ingrese el precio de la Excursión Opcional 1 (Opcional):")
    op2 = st.text_input("Ingrese la Excursión Opcional 2 (Opcional):")
    precio_op2 = st.text_input("Ingrese el precio de la Excursión Opcional 2 (Opcional):")
    
    if st.button("Generar Cartel"):
        archivo_generado = generar_cartel(ciudad, fecha, actividad, hora_encuentro, punto_encuentro, desayuno, nombre_guia, op1, precio_op1, op2, precio_op2, idiomas_seleccionados)
        if archivo_generado.startswith("Error"):
            st.error(archivo_generado)
        else:
            with open(archivo_generado, "rb") as file:
                st.download_button(label="Descargar Cartel", data=file, file_name=os.path.basename(archivo_generado), mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
